from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.shared import Cm

from HuoQuXiuXiRi import get_start, my_work

# Work_month=1
# WorkDates = []
# NoWorkList=[1,6,7,13,14,20,21,27,28]
# for WorkDate in range(0, 31):
#     WorkDates.append(WorkDate + 1)
# for NoWorkDate in NoWorkList:
#     WorkDates.remove(NoWorkDate)
# print("工作日",WorkDates,"工作天数：",len(WorkDates))

Work_month = 1
get_start(2024, Work_month)
WorkDates = my_work
A = len(WorkDates)
QingJiaList = []
for NoWorkDate in QingJiaList:
    WorkDates.remove(NoWorkDate)
B = len(WorkDates)
print("应工作{}天，实际工作{}天".format(A, B))
# 创建文档
document = Document()

for WorkDate in WorkDates:
    p = document.add_paragraph('{}月{}日'.format(Work_month, WorkDate))
    for x in range(0, 2):
        # 添加表格
        tab1 = document.add_table(rows=1, cols=1)  # 添加一个1行1列的空表
        cell = tab1.cell(0, 0)  # 获取某单元格对象（从0开始索引）
        # 在单元格中添加段落，区块
        c_p1 = cell.paragraphs[0]
        c_p1.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 设置单元格内容居中对齐
        c_run1 = c_p1.add_run()
        # 在单元格中添加图片，我的图片是bar.png，图片和py文件在同一个目录下
        try:
            if x == 0:
                c_run1.add_picture(r'工作截图/' + str(WorkDate) + 'a.png', width=Cm(13))
            elif x == 1:
                c_run1.add_picture(r'工作截图/' + str(WorkDate) + 'b.png', width=Cm(13))
        except Exception as e:
            print("----------", WorkDate, "月失败----------")
            print(e)

    p1 = document.add_paragraph(' ')
    run = p1.add_run()
    run.add_break(WD_BREAK.PAGE)

# 保存.docx文档
document.save(str(Work_month) + '月工作截图.docx')
